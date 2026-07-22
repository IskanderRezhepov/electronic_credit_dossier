from __future__ import annotations

from dataclasses import dataclass, field as dataclass_field
from pathlib import Path
from typing import Literal

import hashlib
import json
import re

import fitz
import pytesseract
from docx import Document
from PIL import Image, ImageEnhance, ImageFilter, ImageOps
from pytesseract import Output

from .text_utils import normalize_text

SUPPORTED_EXTENSIONS = {'.pdf', '.png', '.jpg', '.jpeg', '.tif', '.tiff', '.bmp', '.webp', '.docx', '.txt'}


@dataclass
class PageContent:
    page_number: int
    text: str
    extraction_method: Literal['digital', 'ocr', 'hybrid', 'none']
    char_count: int
    quality: float
    variants: dict[str, str] = dataclass_field(default_factory=dict)
    cache_hit: bool = False
    layout_words: list[dict] = dataclass_field(default_factory=list)
    page_width: float | None = None
    page_height: float | None = None


@dataclass
class ReadDocument:
    filename: str
    page_count: int
    source_type: str
    pages: list[PageContent]

    @property
    def full_text(self) -> str:
        return '\n'.join(page.text for page in self.pages)

    @property
    def used_ocr(self) -> bool:
        return any(page.extraction_method in {'ocr', 'hybrid'} for page in self.pages)


class OCRUnavailableError(RuntimeError):
    pass


def read_document(path: Path, *, ocr_languages: str, ocr_dpi: int, min_digital_chars: int, ocr_cache_dir: Path | None = None, analysis_mode: str = 'standard') -> ReadDocument:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f'Неподдерживаемый формат: {suffix}')
    if suffix == '.pdf':
        return _read_pdf(path, ocr_languages, ocr_dpi, min_digital_chars, ocr_cache_dir, analysis_mode)
    if suffix == '.docx':
        return _read_docx(path)
    if suffix == '.txt':
        text = normalize_text(path.read_text(encoding='utf-8', errors='ignore'))
        return ReadDocument(path.name, 1, 'text', [PageContent(1, text, 'digital', len(text), 0.99, {'digital': text})])
    return _read_image(path, ocr_languages, analysis_mode)


def _read_docx(path: Path) -> ReadDocument:
    doc = Document(path)
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            parts.append(' | '.join(cell.text.strip() for cell in row.cells))
    text = normalize_text('\n'.join(parts))
    return ReadDocument(path.name, 1, 'docx', [PageContent(1, text, 'digital', len(text), 0.99, {'digital': text})])



def _digital_layout_words(page: fitz.Page) -> list[dict]:
    words = []
    for item in page.get_text("words") or []:
        if len(item) < 5:
            continue
        x0, y0, x1, y1, value = item[:5]
        value = str(value).strip()
        if not value:
            continue
        words.append({
            "text": value,
            "x0": round(float(x0), 2),
            "y0": round(float(y0), 2),
            "x1": round(float(x1), 2),
            "y1": round(float(y1), 2),
            "confidence": 1.0,
        })
    return words

def _read_pdf(path: Path, languages: str, dpi: int, min_digital_chars: int, cache_dir: Path | None = None, analysis_mode: str = 'standard') -> ReadDocument:
    pages: list[PageContent] = []
    with fitz.open(path) as doc:
        for index, page in enumerate(doc, start=1):
            digital = normalize_text(page.get_text('text') or '')
            digital_words = _digital_layout_words(page)
            page_width = float(page.rect.width)
            page_height = float(page.rect.height)
            if _digital_text_is_usable(digital, min_digital_chars):
                pages.append(PageContent(
                    index, digital, 'digital', len(digital),
                    round(_digital_text_quality(digital), 2), {'digital': digital}, False,
                    digital_words, page_width, page_height,
                ))
                continue

            image = _render_page(page, dpi)
            ocr = _ocr_cached(image, languages, dpi, cache_dir, analysis_mode)
            combined = normalize_text(f'{digital}\n{ocr["preferred"]}' if digital else ocr['preferred'])
            method = 'hybrid' if digital and combined else 'ocr' if combined else 'none'
            pages.append(PageContent(index, combined, method, len(combined), ocr['quality'], ocr['variants'], bool(ocr.get('cache_hit')), ocr.get('layout_words', []), float(image.width), float(image.height)))

        return ReadDocument(path.name, doc.page_count, 'pdf', pages)


def _read_image(path: Path, languages: str, analysis_mode: str = 'standard') -> ReadDocument:
    try:
        image = ImageOps.exif_transpose(Image.open(path)).convert('RGB')
        ocr = _ocr_multimode(image, languages, analysis_mode)
    except Exception as exc:
        raise RuntimeError(f'Не удалось прочитать изображение: {exc}') from exc
    text = normalize_text(ocr['preferred'])
    return ReadDocument(path.name, 1, 'image', [PageContent(1, text, 'ocr', len(text), ocr['quality'], ocr['variants'], bool(ocr.get('cache_hit')), ocr.get('layout_words', []), float(image.width), float(image.height))])


def _digital_text_quality(text: str) -> float:
    """
    Estimate whether the embedded PDF text is genuinely readable.

    Some scanned PDFs contain a broken hidden OCR layer with thousands of
    characters such as ``3aJIOrO...``. Character count alone incorrectly marks
    this as perfect digital text. The score below is deliberately tailored to
    Russian/Kazakh legal documents while still allowing normal Latin names,
    contract numbers and IBANs.
    """
    if not text:
        return 0.0

    letters = [ch for ch in text if ch.isalpha()]
    if not letters:
        return 0.0

    cyrillic = sum(
        "\u0400" <= ch <= "\u052f"
        or ch in "ӘәҒғҚқҢңӨөҰұҮүІіҺһ"
        for ch in letters
    )
    latin = sum(("A" <= ch <= "Z") or ("a" <= ch <= "z") for ch in letters)

    tokens = re.findall(r"\b[\w-]{4,}\b", text, re.UNICODE)
    mixed_alnum = sum(
        bool(re.search(r"[A-Za-zА-Яа-яӘәҒғҚқҢңӨөҰұҮүІіҺһ]", token))
        and bool(re.search(r"\d", token))
        for token in tokens
    )
    mixed_ratio = mixed_alnum / max(len(tokens), 1)

    # Broken OCR frequently contains long Latin-looking transliterations of
    # Cyrillic words. Real banking documents still contain much more Cyrillic
    # prose than Latin prose.
    latin_heavy = latin > cyrillic * 0.75 and cyrillic < len(letters) * 0.55

    common_terms = (
        "ДОГОВОР", "СОГЛАШЕНИ", "СТОРОН", "БИН", "ИИН", "ТЕНГЕ",
        "СЧЕТ", "СЧЁТ", "ЗАЛОГ", "ЛИЗИНГ", "КЕПІЛ", "ШАРТ",
    )
    upper = text.upper()
    language_hits = sum(term in upper for term in common_terms)

    base = min(1.0, len(text) / 1200)
    alphabetic_ratio = len(letters) / max(len(text), 1)
    score = 0.45 * base + 0.35 * min(1.0, alphabetic_ratio / 0.45)
    score += min(0.20, language_hits * 0.035)

    if mixed_ratio > 0.08:
        score -= 0.35
    if latin_heavy:
        score -= 0.30

    return max(0.0, min(1.0, score))


def _digital_text_is_usable(text: str, minimum: int) -> bool:
    if len(text) < minimum:
        return False
    return _digital_text_quality(text) >= 0.62


def _render_page(page: fitz.Page, dpi: int) -> Image.Image:
    zoom = dpi / 72
    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
    return Image.frombytes('RGB', [pix.width, pix.height], pix.samples)


def _preprocess_variants(image: Image.Image) -> dict[str, Image.Image]:
    gray = ImageOps.exif_transpose(image).convert('L')
    gray = ImageOps.autocontrast(gray, cutoff=1)
    contrast = ImageEnhance.Contrast(gray).enhance(1.45).filter(ImageFilter.SHARPEN)
    threshold = contrast.point(lambda p: 255 if p > 175 else 0)
    return {'gray': gray, 'contrast': contrast, 'threshold': threshold}



def _ocr_cached(image: Image.Image, languages: str, dpi: int, cache_dir: Path | None, analysis_mode: str = 'standard') -> dict:
    if cache_dir is None:
        result = _ocr_multimode(image, languages, analysis_mode)
        result['cache_hit'] = False
        return result

    cache_dir.mkdir(parents=True, exist_ok=True)
    digest = hashlib.sha256()
    digest.update(image.tobytes())
    digest.update(f"|{languages}|{dpi}|adaptive-v2|{analysis_mode}".encode("utf-8"))
    cache_path = cache_dir / f"{digest.hexdigest()}.json"

    if cache_path.exists():
        try:
            cached = json.loads(cache_path.read_text(encoding="utf-8"))
            if isinstance(cached, dict) and "preferred" in cached:
                cached["cache_hit"] = True
                return cached
        except (OSError, ValueError, TypeError):
            pass

    result = _ocr_multimode(image, languages, analysis_mode)
    result['cache_hit'] = False
    try:
        cache_path.write_text(json.dumps(result, ensure_ascii=False), encoding="utf-8")
    except OSError:
        pass
    return result



def _ocr_multimode(image: Image.Image, languages: str, analysis_mode: str = 'standard') -> dict:
    """
    Fast adaptive OCR:
    1. one general pass for every page;
    2. a table pass only when the first pass is weak;
    3. sparse/column passes only for difficult pages.

    This is substantially faster than running 3-5 Tesseract passes on every
    page, especially for long scanned contracts.
    """
    prepared = _preprocess_variants(image)
    width, height = image.size
    variants: dict[str, str] = {}
    qualities: list[float] = []

    mode = analysis_mode if analysis_mode in {"fast", "standard", "accurate"} else "standard"

    full_text, full_quality, full_layout = _ocr_with_quality(prepared["contrast"], languages, psm=3)
    variants["full"] = normalize_text(full_text)
    qualities.append(full_quality)

    enough_text = len(variants["full"]) >= (240 if mode == "fast" else 350)
    strong_page = full_quality >= (0.61 if mode == "fast" else 0.69) and enough_text

    # Fast mode intentionally performs one pass unless the page is almost empty.
    run_table = (
        mode == "accurate"
        or (mode == "standard" and not strong_page)
        or (mode == "fast" and len(variants["full"]) < 90)
    )
    if run_table:
        table_text, table_quality, _table_layout = _ocr_with_quality(prepared["threshold"], languages, psm=6)
        variants["table"] = normalize_text(table_text)
        qualities.append(table_quality)

    best_quality = max(qualities or [0.0])
    best_chars = max((len(value) for value in variants.values()), default=0)

    run_sparse = (
        mode == "accurate"
        or (mode == "standard" and (best_quality < 0.60 or best_chars < 220))
        or (mode == "fast" and best_chars < 70)
    )
    if run_sparse:
        sparse_text, sparse_quality, _sparse_layout = _ocr_with_quality(prepared["gray"], languages, psm=11)
        variants["sparse"] = normalize_text(sparse_text)
        qualities.append(sparse_quality)

    # Column OCR is the most expensive pass. It is disabled in fast mode,
    # conditional in standard mode and more permissive in accurate mode.
    column_threshold = 0.74 if mode == "accurate" else 0.66
    if mode != "fast" and width / max(height, 1) > 0.78 and max(qualities or [0.0]) < column_threshold:
        for side, box in {
            "right_column": (int(width * 0.47), 0, width, height),
            "left_column": (0, 0, int(width * 0.53), height),
        }.items():
            crop = prepared["contrast"].crop(box)
            side_text, side_quality, _side_layout = _ocr_with_quality(crop, languages, psm=6)
            variants[side] = normalize_text(side_text)
            qualities.append(side_quality)

    preferred_parts = [
        variants.get("full", ""),
        variants.get("table", ""),
        variants.get("right_column", ""),
        variants.get("left_column", ""),
        variants.get("sparse", ""),
    ]
    preferred = normalize_text("\n".join(part for part in preferred_parts if part))
    return {
        "preferred": preferred,
        "quality": round(max(qualities or [0.0]), 2),
        "variants": variants,
        "layout_words": full_layout,
    }

def _ocr_with_quality(image: Image.Image, languages: str, psm: int) -> tuple[str, float, list[dict]]:
    try:
        config = f'--oem 3 --psm {psm} -c preserve_interword_spaces=1'
        data = pytesseract.image_to_data(image, lang=languages, config=config, output_type=Output.DICT)
        words: list[str] = []
        layout_words: list[dict] = []
        confidences: list[float] = []
        count = len(data.get('text', []))
        for index in range(count):
            value = str(data.get('text', [''])[index]).strip()
            try:
                conf_value = float(data.get('conf', [-1])[index])
            except (TypeError, ValueError, IndexError):
                conf_value = -1
            if not value:
                continue
            words.append(value)
            if conf_value >= 0:
                confidences.append(conf_value)
            left = float(data.get('left', [0])[index])
            top = float(data.get('top', [0])[index])
            width = float(data.get('width', [0])[index])
            height = float(data.get('height', [0])[index])
            layout_words.append({
                'text': value, 'x0': left, 'y0': top,
                'x1': left + width, 'y1': top + height,
                'confidence': round(max(conf_value, 0) / 100, 3),
            })
        quality = (sum(confidences) / len(confidences) / 100) if confidences else 0.0
        return ' '.join(words), quality, layout_words
    except pytesseract.TesseractNotFoundError as exc:
        raise OCRUnavailableError('Tesseract OCR не установлен или не добавлен в PATH.') from exc
